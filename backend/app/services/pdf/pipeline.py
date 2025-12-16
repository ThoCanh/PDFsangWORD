from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from ...core.config import settings
from ...utils.files import safe_filename, which
from .classifier import pdf_has_text_layer, pdf_text_layer_seems_low_quality
from .docx_postprocess import DocxPostprocessError, normalize_docx_page_breaks
from .image_fallback import pdf_to_docx_images
from .ocr import OcrFailedError, run_ocrmypdf
from .adobe_pdf_services_convert import (
    AdobePdfServicesConvertError,
    convert_pdf_to_docx_adobe_pdf_services,
)
from .aspose_words_convert import AsposeWordsConvertError, convert_pdf_to_docx_aspose_words
from .pdf2docx_convert import Pdf2DocxConvertError, convert_pdf_to_docx_pdf2docx
from .pdf_text_docx import PdfTextToDocxError, convert_pdf_text_to_docx


@dataclass(frozen=True)
class PdfToDocxResult:
    docx_path: Path
    mode: str  # "tier-a" | "tier-a-ocr" | "tier-b"
    has_text_layer: bool


class EditableConversionUnavailable(RuntimeError):
    pass


class OcrUnavailableError(RuntimeError):
    """Raised when OCR is requested but OCR toolchain is not available on the server."""
    pass


def _docx_text_len(docx_path: Path) -> int:
    try:
        from docx import Document

        d = Document(str(docx_path))
        total = 0
        for p in d.paragraphs:
            total += len((p.text or "").strip())
        return total
    except Exception:  # noqa: BLE001
        return 0


def _docx_space_ratio(docx_path: Path) -> float:
    try:
        from docx import Document

        d = Document(str(docx_path))
        text = "\n".join((p.text or "") for p in d.paragraphs)
        stripped = text.strip()
        if not stripped:
            return 0.0
        return stripped.count(" ") / max(len(stripped), 1)
    except Exception:  # noqa: BLE001
        return 0.0


def _docx_looks_mojibake(docx_path: Path) -> bool:
    try:
        from docx import Document

        d = Document(str(docx_path))
        text = "\n".join((p.text or "") for p in d.paragraphs)
        # Simple heuristics: presence of replacement chars or common mojibake markers
        if "\ufffd" in text:
            return True
        # Common UTF-8 mis-decode artifacts for Vietnamese can contain 'Ã' sequences
        if "Ã" in text:
            return True
        # Some PDFs yield odd sequences like 'D<' from font mapping issues
        if "<" in text and any(c.isalpha() for c in text.split("<")[0][-2:]):
            return True
        return False
    except Exception:  # noqa: BLE001
        return False


def _pdf_text_len(pdf_path: Path, max_pages: int) -> int:
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(pdf_path))
        try:
            total_pages = doc.page_count
            pages_to_count = total_pages
            if max_pages and int(max_pages) > 0:
                pages_to_count = min(total_pages, int(max_pages))
            total = 0
            for i in range(pages_to_count):
                total += len((doc.load_page(i).get_text("text") or "").strip())
            return total
        finally:
            doc.close()
    except Exception:  # noqa: BLE001
        return 0


def _pdf_text_looks_mojibake(pdf_path: Path, max_pages: int = 2) -> bool:
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(pdf_path))
        try:
            pages_to_check = min(doc.page_count, max_pages)
            text = []
            for i in range(pages_to_check):
                text.append(doc.load_page(i).get_text("text") or "")
            joined = "\n".join(text)
            # Replacement character is a definite sign of decoding issues.
            if "\ufffd" in joined:
                return True
            # Detect common UTF-8 -> Windows-1252 mojibake patterns like 'Ã' or 'Â' followed by continuation-byte range
            for i, ch in enumerate(joined[:-1]):
                o = ord(ch)
                if o in (0x00C2, 0x00C3):
                    next_o = ord(joined[i+1])
                    if 128 <= next_o <= 191:
                        return True
            # Check for suspicious '<' sequences e.g., 'D<' often indicates garbled spacing
            if "<" in joined and any(c.isalpha() for c in joined.split("<")[0][-2:]):
                return True
            return False
        finally:
            doc.close()
    except Exception:  # noqa: BLE001
        return False


def convert_pdf_to_docx_pipeline(*, pdf_path: Path, work_dir: Path, prefer_tier_a: bool = False, force_ocr: bool = False) -> PdfToDocxResult:
    """Professional PDF→DOCX pipeline.

    Tier A: Adobe PDF Services API (when configured), then Aspose.Words, then pdf2docx
    Tier B: Image fallback (max visual fidelity)
    """

    if not pdf_path.exists():
        raise FileNotFoundError(str(pdf_path))

    has_text_initial = pdf_has_text_layer(pdf_path)
    has_text = has_text_initial

    # Some PDFs contain a "text layer" that is effectively unusable for high-quality
    # DOCX conversion (often OCR output with missing spaces / poor glyph mapping).
    # For Vietnamese, this frequently shows up as words glued together.
    if has_text and settings.ocr_enabled and pdf_text_layer_seems_low_quality(pdf_path):
        has_text = False

    out_dir = work_dir / "tier-a"
    adobe_error: str | None = None
    aspose_error: str | None = None
    pdf2docx_error: str | None = None
    ocr_error: str | None = None
    postprocess_error: str | None = None
    fallback_error: str | None = None

    adobe_enabled = bool(settings.adobe_client_id and settings.adobe_client_secret)

    # OCR-first flow: run OCR when explicitly requested (force_ocr) or when in auto mode the
    # PDF appears scanned and prefer_tier_a is not set. OCR is performed locally via OCRmyPDF + Tesseract
    # and will produce a searchable PDF. We intentionally avoid image cleaning or aggressive processing
    # so stamps/con dấu remain intact (we do NOT use --clean).
    if (force_ocr or ((not prefer_tier_a) and (not has_text))) and settings.ocr_enabled:
        ocrmypdf = which("ocrmypdf", settings.ocrmypdf_path)
        if not ocrmypdf:
            # Explicitly fail so caller can inform admin to install OCR tools for the OCR-first path
            raise OcrUnavailableError(
                "OCR chưa được cài đặt trên máy chủ. Vui lòng cài đặt ocrmypdf và Tesseract để sử dụng OCR-first khi cần."
            )
        try:
            ocr_out = work_dir / "ocr" / "searchable.pdf"
            # Use OCRmyPDF to create a searchable PDF while preserving original images/graphics
            run_ocrmypdf(
                input_pdf=pdf_path,
                output_pdf=ocr_out,
                ocrmypdf_path=ocrmypdf,
                lang=settings.ocr_lang,
                timeout_sec=settings.ocr_timeout_sec,
                extra_path=settings.tesseract_path,
            )
            # Quick check: ensure OCR result doesn't look like Mojibake (wrong encoding)
            if _pdf_text_looks_mojibake(ocr_out, max_pages=2):
                raise EditableConversionUnavailable(
                    "OCR cục bộ tạo lớp text nhưng phát hiện Mojibake (lỗi font). Hãy đảm bảo Tesseract đã có traineddata 'vie' và TESSDATA_PREFIX/TESSERACT_PATH được cấu hình đúng."
                )
            # Use OCRed PDF for subsequent Tier A conversions
            pdf_path = ocr_out
            has_text = True
        except EditableConversionUnavailable:
            raise
        except Exception as e:
            # Surface OCR failure explicitly
            raise OcrUnavailableError(f"OCR failed: {e}") from e

    # Tier A (Adobe PDF Services API preferred when configured)
    if adobe_enabled:
        try:
            def _run_adobe(*, ocr_lang: str | None) -> tuple[Path, str]:
                mode_local = "tier-a-adobe" if not ocr_lang else "tier-a-adobe-ocr"
                r = convert_pdf_to_docx_adobe_pdf_services(
                    pdf_path=pdf_path,
                    out_dir=out_dir,
                    base_url=settings.adobe_base_url,
                    client_id=str(settings.adobe_client_id),
                    client_secret=str(settings.adobe_client_secret),
                    job_timeout_sec=settings.adobe_job_timeout_sec,
                    poll_interval_ms=settings.adobe_poll_interval_ms,
                    ocr_lang=ocr_lang,
                )
                return r.docx_path, mode_local

            # Decide whether to request Adobe OCR. If we ran local OCR (force_ocr), we DO NOT request Adobe OCR
            # so Adobe will use our searchable PDF layer. If user explicitly asked Tier A, we still may request Adobe OCR
            # if configured (prefer_tier_a semantics). For scanned PDFs where we did local OCR, set ocr_lang=None.
            ocr_lang: str | None = None
            if prefer_tier_a:
                ocr_lang = settings.adobe_ocr_lang
            elif (not has_text) and settings.ocr_enabled:
                ocr_lang = settings.adobe_ocr_lang

            docx_path, mode = _run_adobe(ocr_lang=ocr_lang)

            # If we explicitly ran OCR locally (force_ocr), prefer to mark the mode as "tier-a-ocr"
            if force_ocr:
                mode = "tier-a-ocr"

            # If we did NOT OCR (because PDF appeared to have text), but the resulting
            # DOCX still looks like it has no spaces, retry once with Adobe OCR.
            if not prefer_tier_a and has_text and settings.ocr_enabled:
                space_ratio = _docx_space_ratio(docx_path)
                if space_ratio > 0 and space_ratio < 0.01:
                    docx_path, mode = _run_adobe(ocr_lang=settings.adobe_ocr_lang)
                    mode = "tier-a-adobe-ocr-retry"

            adobe_result_docx = docx_path

            try:
                normalize_docx_page_breaks(docx_path=adobe_result_docx, aggressive=False)
            except DocxPostprocessError as e:
                postprocess_error = str(e)

            # If we forced local OCR before sending to Adobe, detect any mojibake patterns.
            # Try Aspose fallback if Adobe produced mojibake after our OCR, otherwise fail with helpful guidance.
            if force_ocr and _docx_looks_mojibake(adobe_result_docx):
                # Attempt Aspose fallback using the searchable PDF (pdf_path should point to OCRed PDF)
                try:
                    aspose_fallback_dir = out_dir / "aspose-after-ocr"
                    aspose_fallback = convert_pdf_to_docx_aspose_words(pdf_path=pdf_path, out_dir=aspose_fallback_dir)
                    try:
                        normalize_docx_page_breaks(docx_path=aspose_fallback.docx_path, aggressive=False)
                    except DocxPostprocessError as e:
                        postprocess_error = str(e)

                    # If Aspose also looks mojibake, fail explicitly
                    if _docx_looks_mojibake(aspose_fallback.docx_path):
                        raise EditableConversionUnavailable(
                            "Cả Adobe và Aspose trên kết quả OCR cục bộ đều chứa dấu hiệu Mojibake. Vui lòng kiểm tra rằng Tesseract đã dùng traineddata 'vie' và TESSDATA_PREFIX/TESSERACT_PATH đúng."
                        )

                    # Otherwise, return Aspose result as a fallback (less ideal layout but preserves text)
                    return PdfToDocxResult(
                        docx_path=aspose_fallback.docx_path,
                        mode="aspose-after-ocr",
                        has_text_layer=True,
                    )
                except Exception as e_fallback:
                    # Surface a combined error explaining both failures
                    raise EditableConversionUnavailable(
                        "Sau khi OCR cục bộ (local) -> Adobe, kết quả chứa dấu hiệu Mojibake; thử fallback bằng Aspose cũng thất bại: "
                        + str(e_fallback)
                    ) from e_fallback

            # Missing-content guard for text-layer PDFs
            if has_text:
                pdf_len = _pdf_text_len(pdf_path, settings.max_pages)
                docx_len = _docx_text_len(adobe_result_docx)
                if pdf_len >= 300 and docx_len < int(pdf_len * 0.15):
                    try:
                        docx_text = convert_pdf_text_to_docx(
                            pdf_path=pdf_path,
                            out_dir=out_dir / "text-fallback",
                            max_pages=settings.max_pages,
                        )
                        return PdfToDocxResult(
                            docx_path=docx_text,
                            mode="tier-a-text",
                            has_text_layer=has_text_initial,
                        )
                    except PdfTextToDocxError as e:
                        fallback_error = str(e)

            return PdfToDocxResult(
                docx_path=adobe_result_docx,
                mode=mode,
                has_text_layer=has_text_initial,
            )
        except AdobePdfServicesConvertError as e:
            adobe_error = str(e)

    # Tier A (Aspose.Words preferred):
    # - If PDF already has a text layer, convert directly.
    # - If scanned, OCR to searchable PDF then convert.
    try:
        aspose_result = convert_pdf_to_docx_aspose_words(pdf_path=pdf_path, out_dir=out_dir)
        try:
            normalize_docx_page_breaks(docx_path=aspose_result.docx_path, aggressive=False)
        except DocxPostprocessError as e:
            postprocess_error = str(e)

        # If the PDF has selectable text but the produced DOCX contains far less text,
        # fall back to a text-only DOCX to avoid "mất nội dung".
        if has_text:
            pdf_len = _pdf_text_len(pdf_path, settings.max_pages)
            docx_len = _docx_text_len(aspose_result.docx_path)
            if pdf_len >= 300 and docx_len < int(pdf_len * 0.15):
                try:
                    docx_text = convert_pdf_text_to_docx(
                        pdf_path=pdf_path,
                        out_dir=out_dir / "text-fallback",
                        max_pages=settings.max_pages,
                    )
                    return PdfToDocxResult(
                        docx_path=docx_text,
                        mode="tier-a-text",
                        has_text_layer=has_text,
                    )
                except PdfTextToDocxError as e:
                    fallback_error = str(e)

        return PdfToDocxResult(
            docx_path=aspose_result.docx_path,
            mode="tier-a",
            has_text_layer=has_text,
        )
    except AsposeWordsConvertError as e:
        aspose_error = str(e)

    # Fallback: pdf2docx (still useful when Aspose isn't installed/working)
    try:
        docx_result = convert_pdf_to_docx_pdf2docx(
            pdf_path=pdf_path,
            out_dir=out_dir,
            max_pages=settings.max_pages,
        )
        try:
            normalize_docx_page_breaks(docx_path=docx_result.docx_path, aggressive=True)
        except DocxPostprocessError as e:
            postprocess_error = str(e)
        if has_text:
            pdf_len = _pdf_text_len(pdf_path, settings.max_pages)
            docx_len = _docx_text_len(docx_result.docx_path)
            if pdf_len >= 300 and docx_len < int(pdf_len * 0.15):
                try:
                    docx_text = convert_pdf_text_to_docx(
                        pdf_path=pdf_path,
                        out_dir=out_dir / "text-fallback",
                        max_pages=settings.max_pages,
                    )
                    return PdfToDocxResult(
                        docx_path=docx_text,
                        mode="tier-a-text",
                        has_text_layer=has_text,
                    )
                except PdfTextToDocxError as e:
                    fallback_error = str(e)

        return PdfToDocxResult(
            docx_path=docx_result.docx_path,
            mode="tier-a",
            has_text_layer=has_text,
        )
    except Pdf2DocxConvertError as e:
        pdf2docx_error = str(e)

    if (not has_text) and settings.ocr_enabled:
        ocrmypdf = which("ocrmypdf", settings.ocrmypdf_path)
        if ocrmypdf:
            ocr_out = work_dir / "ocr" / "searchable.pdf"
            try:
                run_ocrmypdf(
                    input_pdf=pdf_path,
                    output_pdf=ocr_out,
                    ocrmypdf_path=ocrmypdf,
                    lang=settings.ocr_lang,
                    timeout_sec=settings.ocr_timeout_sec,
                    extra_path=settings.tesseract_path,
                )
                has_text_after = pdf_has_text_layer(ocr_out)
                # Prefer Aspose after OCR
                try:
                    aspose2 = convert_pdf_to_docx_aspose_words(pdf_path=ocr_out, out_dir=out_dir)
                    try:
                        normalize_docx_page_breaks(docx_path=aspose2.docx_path, aggressive=False)
                    except DocxPostprocessError as e:
                        postprocess_error = str(e)
                    return PdfToDocxResult(
                        docx_path=aspose2.docx_path,
                        mode="tier-a-ocr",
                        has_text_layer=has_text_after,
                    )
                except AsposeWordsConvertError as e_aspose_ocr:
                    aspose_error = str(e_aspose_ocr)

                # Fallback to pdf2docx after OCR
                docx2_result = convert_pdf_to_docx_pdf2docx(
                    pdf_path=ocr_out,
                    out_dir=out_dir,
                    max_pages=settings.max_pages,
                )
                try:
                    normalize_docx_page_breaks(docx_path=docx2_result.docx_path, aggressive=True)
                except DocxPostprocessError as e:
                    postprocess_error = str(e)
                return PdfToDocxResult(
                    docx_path=docx2_result.docx_path,
                    mode="tier-a-ocr",
                    has_text_layer=has_text_after,
                )
            except (OcrFailedError, Pdf2DocxConvertError) as e2:
                ocr_error = str(e2)

    if settings.prefer_editable:
        raise EditableConversionUnavailable(
            "Unable to produce editable DOCX. "
            f"has_text_layer={has_text}. "
            f"adobe_error={adobe_error or 'n/a'}. "
            f"aspose_error={aspose_error or 'n/a'}. "
            f"pdf2docx_error={pdf2docx_error or 'n/a'}. "
            f"ocr_error={ocr_error or 'n/a'}. "
            f"postprocess_error={postprocess_error or 'n/a'}. "
            f"text_fallback_error={fallback_error or 'n/a'}. "
            "If this PDF is scanned, ensure OCRmyPDF + Tesseract (and Ghostscript on Windows) are installed, or disable PREFER_EDITABLE to allow image fallback."
        )

    # Tier B (Image fallback)
    stem = safe_filename(pdf_path.stem, fallback="document")
    out_docx = work_dir / "tier-b" / f"{stem}.docx"
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    docx = pdf_to_docx_images(
        pdf_path=pdf_path,
        out_docx=out_docx,
        dpi=settings.pdf_image_dpi,
        max_pages=settings.max_pages,
    )
    return PdfToDocxResult(docx_path=docx, mode="tier-b", has_text_layer=has_text)

from __future__ import annotations

import io
from pathlib import Path


class ImageFallbackError(RuntimeError):
    pass


def pdf_to_docx_images(
    *,
    pdf_path: Path,
    out_docx: Path,
    dpi: int,
    max_pages: int,
) -> Path:
    """Tier B: render each PDF page to an image and embed into DOCX.

    This maximizes visual fidelity ("looks identical") but results are not fully editable.
    """

    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Emu, Inches

    if dpi < 72:
        dpi = 72

    doc = fitz.open(str(pdf_path))
    try:
        page_count = doc.page_count
        if page_count == 0:
            raise ImageFallbackError("PDF has no pages")
        if page_count > max_pages:
            raise ImageFallbackError(f"PDF has too many pages ({page_count}); max is {max_pages}")

        word_doc = Document()
        section = word_doc.sections[0]

        # Reduce margins to preserve page fit.
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.25)

        usable_width = section.page_width - section.left_margin - section.right_margin
        usable_height = section.page_height - section.top_margin - section.bottom_margin

        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)

        for idx in range(page_count):
            page = doc.load_page(idx)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")

            p = word_doc.add_paragraph()
            p.paragraph_format.space_before = 0
            p.paragraph_format.space_after = 0

            run = p.add_run()
            # Scale to fit BOTH width and height to avoid overflow that can create blank pages.
            max_w_emu = int(usable_width)
            max_h_emu = int(usable_height)
            img_w = max(pix.width, 1)
            img_h = max(pix.height, 1)

            # Start by fitting width.
            w_emu = max_w_emu
            h_emu = int(w_emu * (img_h / img_w))

            # If too tall, fit height instead.
            if h_emu > max_h_emu:
                h_emu = max_h_emu
                w_emu = int(h_emu * (img_w / img_h))

            run.add_picture(io.BytesIO(img_bytes), width=Emu(w_emu), height=Emu(h_emu))

            # Intentionally avoid explicit page breaks; Word will paginate based on content height.

        out_docx.parent.mkdir(parents=True, exist_ok=True)
        word_doc.save(str(out_docx))
        if out_docx.stat().st_size == 0:
            raise ImageFallbackError("Generated DOCX is empty")
        return out_docx
    finally:
        doc.close()

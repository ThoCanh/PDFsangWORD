from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from ...utils.files import safe_filename


class AsposeWordsConvertError(RuntimeError):
    pass


@dataclass(frozen=True)
class AsposeWordsConvertResult:
    docx_path: Path


def remove_aspose_watermark(docx_path: str | Path) -> None:
    """Xóa watermark của Aspose.Words từ file DOCX.
    
    Watermark có thể xuất hiện dưới dạng:
    - "Created with an evaluation copy of Aspose.Words..."
    - "Evaluation Only"
    - Các phần khác của text watermark
    """
    try:
        from docx import Document
        
        doc = Document(str(docx_path))
        paragraphs_to_delete = []
        
        # Tìm các paragraph chứa watermark
        watermark_keywords = [
            "Evaluation Only",
            "Aspose.Words",
            "evaluation copy",
            "temporary-license",
            "products.aspose.com",
            "Created with an evaluation copy",
        ]
        
        for para in doc.paragraphs:
            para_text = para.text.lower()
            # Kiểm tra nếu paragraph chứa bất kỳ keyword nào
            if any(keyword.lower() in para_text for keyword in watermark_keywords):
                paragraphs_to_delete.append(para)
        
        # Xóa các paragraph chứa watermark
        for para in paragraphs_to_delete:
            p = para._element
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)
                para._p = None
                para._element = None
        
        # Lưu file sau khi xóa watermark
        doc.save(str(docx_path))
    except Exception:  # noqa: BLE001
        # Nếu có lỗi khi xóa watermark, không làm gián đoạn quá trình chuyển đổi
        pass


def convert_pdf_to_docx_aspose_words(*, pdf_path: Path, out_dir: Path) -> AsposeWordsConvertResult:
    """Convert PDF → DOCX using Aspose.Words.

    Notes:
    - Aspose.Words is commercial. Without a license it may add evaluation watermarks.
    - It does NOT perform OCR by itself; for scanned PDFs, run OCR first.
    - Watermarks are automatically removed after conversion.
    - Fonts are embedded in the output DOCX for better compatibility.
    """

    if not pdf_path.exists():
        raise FileNotFoundError(str(pdf_path))

    out_dir.mkdir(parents=True, exist_ok=True)
    stem = safe_filename(pdf_path.stem, fallback="document")
    out_docx = out_dir / f"{stem}.docx"

    try:
        import aspose.words as aw

        # Load PDF với options
        load_options = aw.loading.PdfLoadOptions()
        load_options.skip_pdf_images = False
        doc = aw.Document(str(pdf_path), load_options)

        # Cấu hình nhúng font (embed fonts) để đảm bảo hiển thị đúng
        font_infos = doc.font_infos
        font_infos.embed_true_type_fonts = True
        font_infos.embed_system_fonts = True
        font_infos.save_subset_fonts = True  # Chỉ nhúng các ký tự được dùng để giảm dung lượng

        # Tối ưu layout
        doc.update_page_layout()

        # Lưu file DOCX
        doc.save(str(out_docx), aw.SaveFormat.DOCX)

        if not out_docx.exists():
            raise AsposeWordsConvertError("Aspose.Words did not produce a DOCX output")

        # Xóa watermark tự động sau khi chuyển đổi
        remove_aspose_watermark(out_docx)

        return AsposeWordsConvertResult(docx_path=out_docx)

    except AsposeWordsConvertError:
        raise
    except Exception as e:  # noqa: BLE001
        raise AsposeWordsConvertError(str(e)) from e
